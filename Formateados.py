from docx import Document
from docx.shared import Pt
import re

def formatear_documento(nombre_archivo: str, nombre_salida: str):
    """
    Formatea un documento de Word agregando un salto de línea antes de cada artículo y poniendo
    en negrita la palabra "Artículo" junto con su número.
    Además, establece la fuente a Times New Roman, tamaño 12 para todo el texto.
    
    Argumentos:
    - nombre_archivo (str): Ruta del archivo de Word a procesar.
    - nombre_salida (str): Ruta donde se guardará el documento formateado.
    """
    doc = Document(nombre_archivo)
    nuevo_doc = Document()
    
    # Aplicar formato base al documento
    estilo = nuevo_doc.styles['Normal']
    estilo.font.name = 'Times New Roman'
    estilo.font.size = Pt(12)
    
    patron_articulo = re.compile(r'^(Artículo \d+)(.-)', re.MULTILINE)
    
    for parrafo in doc.paragraphs:
        texto = parrafo.text.strip()
        
        if texto.startswith("Artículo"):
            nuevo_doc.add_paragraph("")  # Agrega un espacio antes de cada artículo
            nuevo_parrafo = nuevo_doc.add_paragraph()
            nuevo_parrafo.style = nuevo_doc.styles['Normal']
            
            match = patron_articulo.match(texto)
            if match:
                inicio_negrita = match.group(1)  # "Artículo X"
                resto_texto = texto[len(inicio_negrita):]  # El resto del texto sin el inicio
                
                run = nuevo_parrafo.add_run(inicio_negrita)
                run.bold = True  # Poner en negrita "Artículo X"
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                nuevo_parrafo.add_run(resto_texto)  # Agregar el resto del texto normal
            else:
                run = nuevo_parrafo.add_run(texto)  # Si no coincide, agregar el texto como está
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        else:
            nuevo_parrafo = nuevo_doc.add_paragraph(texto)
            nuevo_parrafo.style = nuevo_doc.styles['Normal']
    
    nuevo_doc.save(nombre_salida)

# Uso del script (reemplaza los nombres de archivo con los tuyos)
nombre_archivo = "archivo_original
nombre_salida = "archivo_formateado"
formatear_documento(nombre_archivo, nombre_salida)
